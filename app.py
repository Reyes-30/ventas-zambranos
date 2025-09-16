"""
Sistema de Análisis de Ventas Zambranos
=========================================

Aplicación web moderna para análisis avanzado de datos de ventas con capacidades de 
Machine Learning, visualizaciones interactivas y sistema de exportación profesional.

La aplicación está construida con una arquitectura modular que incluye:
- Análisis estadístico descriptivo e inferencial
- Algoritmos de Machine Learning (PCA, K-Means)
- Visualizaciones interactivas con Plotly y estáticas con Matplotlib
- Sistema de exportación en múltiples formatos
- Tour interactivo guiado para usuarios
- Manejo seguro de archivos con validación
- Sistema de logging estructurado

Desarrollado con Streamlit como framework web principal.

Author: Equipo de Desarrollo Zambranos
Version: 2.0
Last Updated: Septiembre 2025
"""

import io
import csv
import sys
from typing import Optional, Tuple
import base64
import zipfile
from io import BytesIO
from docx import Document

import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.cluster import KMeans
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio

import streamlit as st
import logging
from pathlib import Path
from src.app.io_utils import (
    read_dataframe as safe_read_dataframe,
    validate_schema as safe_validate_schema,
    coerce_numeric as safe_coerce_numeric,
    save_upload_safely,
    UPLOADS_DIR,
)
from src.app.processing import (
    MESES_ORDEN, VARIABLES_NUM, resumen_metricas, series_temporales,
    categoria_agg, run_pca, run_kmeans
)
from src.app.exceptions import AppError, ValidationError, FileIOError, ProcessingError
from src.app.tour_interactivo import mostrar_tour_interactivo, mostrar_ayuda_contextual

# Configuración del Sistema de Logging
# ====================================
# 
# El sistema de logging está configurado para registrar:
# - Acciones del usuario (cargas de archivos, análisis ejecutados)
# - Errores y excepciones con contexto completo
# - Métricas de rendimiento y uso de la aplicación
# - Logs tanto en archivo como en consola para debugging

LOG_DIR = Path(__file__).resolve().parent / "logs"
LOG_DIR.mkdir(exist_ok=True)  # Crear directorio de logs si no existe

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_DIR / "app.log", encoding="utf-8"),  # Archivo persistente
        logging.StreamHandler()  # Consola para desarrollo
    ]
)
logger = logging.getLogger(__name__)

# Configuración de Página con menú personalizado
# ================================================
# Configura la página principal de Streamlit con opciones personalizadas
# y oculta elementos no relevantes para usuarios finales

st.set_page_config(
    page_title="📊 Sistema de Análisis de Ventas - Zambranos",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/tu-repo/docs/manual_usuario.md',  # Link a tu documentación
        'Report a bug': 'https://github.com/tu-repo/issues',  # Link a issues de GitHub
        'About': """
        # Sistema de Análisis de Ventas Zambranos v2.0
        
        **Aplicación web moderna para análisis avanzado de datos de ventas**
        
        ### Características:
        - 📊 Dashboard ejecutivo con métricas clave
        - 📈 Análisis temporal y por categorías  
        - 🤖 Machine Learning (PCA, K-Means)
        - 📥 Exportación profesional en múltiples formatos
        - 🎓 Tour interactivo para nuevos usuarios
        
        ### Desarrollado por:
        Equipo de Desarrollo Zambranos
        
        ### Tecnologías:
        Streamlit • Pandas • Plotly • Scikit-learn
        """
    }
)

# CSS personalizado para mejorar la interfaz y ocultar elementos no relevantes
st.markdown("""
<style>
    /* Ocultar menú hamburguesa de Streamlit */
    #MainMenu {visibility: hidden;}
    
    /* Ocultar footer de Streamlit */
    footer {visibility: hidden;}
    
    /* Mantener visible el header para permitir mostrar/ocultar la sidebar */
    /* (antes estaba oculto y el usuario no podía reabrir la barra lateral) */
    /* header {visibility: hidden;} */
    
    /* Estilo del header principal */
    .main-header {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    /* Estilo de las métricas */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        border-left: 4px solid #2a5298;
        margin-bottom: 1rem;
    }
    
    /* Estilo de las secciones */
    .section-header {
        background: linear-gradient(90deg, #f8f9fa 0%, #e9ecef 100%);
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #28a745;
        margin: 1.5rem 0 1rem 0;
    }
    
    /* Botones personalizados */
    .stButton > button {
        background: linear-gradient(90deg, #28a745 0%, #20c997 100%);
        color: white;
        border: none;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
    }
    
    /* Sidebar personalizada */
    .css-1d391kg {
        background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
    }
    
    /* Alertas personalizadas */
    .success-alert {
        background: linear-gradient(90deg, #d4edda 0%, #c3e6cb 100%);
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    
    .warning-alert {
        background: linear-gradient(90deg, #fff3cd 0%, #ffeaa7 100%);
        border: 1px solid #ffeaa7;
        color: #856404;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    
    /* Tabs personalizadas */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: linear-gradient(90deg, #6c757d 0%, #495057 100%);
        color: white;
        border-radius: 8px 8px 0 0;
        padding: 0.5rem 1rem;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #007bff 0%, #0056b3 100%);
    }
</style>
""", unsafe_allow_html=True)

sns.set_theme(style="whitegrid", palette="viridis")
plt.rcParams["figure.figsize"] = (12, 6)
plt.rcParams["font.size"] = 10

COLUMNAS_REQUERIDAS = set(["Mes", "Categoría", "Cantidad Vendida", "Ingreso Total", "ISV", "Utilidad Bruta"]) | set(VARIABLES_NUM)

# ------------------------------
# Funciones de utilidad para la UI
# ------------------------------

def crear_tarjeta_metrica(titulo: str, valor: str, delta: str = None, delta_color: str = "normal"):
    """Crea una tarjeta de métrica personalizada"""
    if delta:
        st.metric(
            label=titulo,
            value=valor,
            delta=delta,
            delta_color=delta_color
        )
    else:
        st.metric(label=titulo, value=valor)

def mostrar_header_seccion(titulo: str, descripcion: str = ""):
    """Muestra un header de sección estilizado"""
    st.markdown(f"""
    <div class="section-header">
        <h3 style="margin: 0; color: #2c3e50;">{titulo}</h3>
        {f'<p style="margin: 0.5rem 0 0 0; color: #6c757d; font-size: 0.9rem;">{descripcion}</p>' if descripcion else ''}
    </div>
    """, unsafe_allow_html=True)

def crear_alerta_personalizada(mensaje: str, tipo: str = "info"):
    """Crea alertas personalizadas"""
    clase = "success-alert" if tipo == "success" else "warning-alert"
    st.markdown(f'<div class="{clase}">{mensaje}</div>', unsafe_allow_html=True)

def generar_download_link(fig, filename: str, format: str = "png"):
    """Genera un enlace de descarga para gráficas"""
    img_buffer = io.BytesIO()
    fig.savefig(img_buffer, format=format, bbox_inches='tight', dpi=300)
    img_buffer.seek(0)
    b64 = base64.b64encode(img_buffer.read()).decode()
    
    return f'<a href="data:image/{format};base64,{b64}" download="{filename}.{format}" style="text-decoration: none;"><button style="background: #007bff; color: white; border: none; padding: 0.5rem 1rem; border-radius: 5px; cursor: pointer;">📥 Descargar {format.upper()}</button></a>'

# ---- Registro centralizado de gráficas para exportación masiva ----
if "_assets" not in st.session_state:
    st.session_state["_assets"] = {}

def _record_matplotlib(fig, name: str):
    try:
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
        buf.seek(0)
        st.session_state["_assets"][f"{name}.png"] = buf.read()
    except Exception:
        logger.exception("Fallo guardando figura matplotlib %s", name)

def _record_plotly(fig, name: str):
    try:
        img_bytes = pio.to_image(fig, format="png", scale=2)  # requiere kaleido
        st.session_state["_assets"][f"{name}.png"] = img_bytes
    except Exception:
        logger.exception("Fallo guardando figura plotly %s", name)

def exportar_dataframe_csv(df: pd.DataFrame, nombre: str) -> None:
    csv_bytes = df.to_csv(index=False).encode('utf-8')
    st.download_button(
        label=f"📥 Descargar {nombre}.csv",
        data=csv_bytes,
        file_name=f"{nombre}.csv",
        mime="text/csv"
    )

def exportar_zip(diccionario_archivos: dict[str, bytes], nombre_zip: str):
    mem = BytesIO()
    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for ruta, contenido in diccionario_archivos.items():
            zf.writestr(ruta, contenido)
    mem.seek(0)
    st.download_button(
        label=f"📦 Descargar {nombre_zip}",
        data=mem,
        file_name=nombre_zip,
        mime="application/zip"
    )

def generar_reporte_docx(metricas: dict, notas: str = "") -> bytes:
    doc = Document()
    doc.add_heading('Reporte de Análisis de Ventas - Zambranos', level=1)
    doc.add_paragraph('Resumen de métricas clave:')
    if metricas:
        for k, v in metricas.items():
            doc.add_paragraph(f"- {k}: {v}")
    if notas:
        doc.add_heading('Notas', level=2)
        doc.add_paragraph(notas)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ------------------------------
# Utilidades de carga de datos
# ------------------------------

def _sniff_delimiter(sample: bytes) -> Optional[str]:
    try:
        dialect = csv.Sniffer().sniff(sample.decode("utf-8", errors="ignore"))
        return dialect.delimiter
    except Exception:
        return None


@st.cache_data(show_spinner=False)
def load_local_fallback() -> Tuple[Optional[pd.DataFrame], str]:
    for name in ("datosExcel.csv", "datosExcel.xlsx"):
        try:
            df = safe_read_dataframe(str(name))
            return df, name
        except Exception:
            continue
    return None, ""


@st.cache_data(show_spinner=False)
def load_dataframe(file: Optional[io.BytesIO], file_name: Optional[str], delimiter_choice: str) -> Tuple[Optional[pd.DataFrame], str]:
    """
    Carga y procesa un archivo de datos CSV o Excel con validación y manejo de errores.

    Esta función es el punto de entrada principal para la carga de datos en la aplicación.
    Maneja tanto archivos subidos directamente como archivos de respaldo locales.

    Args:
        file (Optional[io.BytesIO]): Objeto de archivo en memoria (uploaded file)
        file_name (Optional[str]): Nombre del archivo original
        delimiter_choice (str): Opción de delimitador para archivos CSV:
                               "Auto", "Coma (,)", "Punto y coma (;)", "Tab (\t)"

    Returns:
        Tuple[Optional[pd.DataFrame], str]: 
            - DataFrame procesado y validado, None si falla la carga
            - Nombre/origen del archivo cargado para logging

    Raises:
        DataValidationError: Si el archivo no cumple con el esquema requerido
        FileProcessingError: Si hay problemas leyendo el archivo

    Examples:
        >>> df, origen = load_dataframe(uploaded_file, "ventas.csv", "Auto")
        >>> if df is not None:
        >>>     print(f"Cargado: {origen}, filas: {len(df)}")

    Note:
        La función implementa fallback automático a datos locales si falla
        la carga del archivo subido, garantizando funcionamiento continuo.
    """
    if file is not None and file_name:
        name_lower = file_name.lower()
        try:
            if name_lower.endswith(".csv"):
                # Determinar delimitador según selección del usuario
                if delimiter_choice == "Auto":
                    sep = None  # Detección automática en safe_read_dataframe
                elif delimiter_choice == "Coma (,)":
                    sep = ","
                elif delimiter_choice == "Punto y coma (;)":
                    sep = ";"
                else:  # Tab
                    sep = "\t"
                
                # Usar función segura de lectura con validación integrada
                df = safe_read_dataframe(file.getvalue(), delimiter=sep)
                logger.info(f"CSV cargado exitosamente: {file_name}, filas: {len(df)}")
                return df, file_name
                
            elif name_lower.endswith(".xlsx") or name_lower.endswith(".xls"):
                # Cargar Excel usando pandas engine por defecto
                df = safe_read_dataframe(file.getvalue())
                logger.info(f"Excel cargado exitosamente: {file_name}, filas: {len(df)}")
                return df, file_name
                
        except Exception as e:
            logger.error(f"Error cargando archivo {file_name}: {str(e)}")
            st.error(f"No se pudo leer el archivo subido: {e}")
            return None, ""

    # Fallback automático a archivos locales si no hay archivo subido
    # Esto garantiza que la aplicación pueda funcionar con datos de demostración
    df_local, origin = load_local_fallback()
    if df_local is not None:
        st.info(f"Usando datos locales de fallback: {origin}")
        logger.info(f"Fallback a datos locales: {origin}")
    else:
        st.warning("Sube un archivo CSV/XLSX con las columnas requeridas para continuar.")
        logger.warning("No hay datos disponibles - requiere upload de usuario")
    
    return df_local, origin


def validar_columnas(df: pd.DataFrame) -> Tuple[bool, set]:
    """
    Valida que el DataFrame contenga todas las columnas requeridas para el análisis.

    Utiliza la función de validación segura del módulo io_utils para verificar
    que el esquema de datos cumple con los requisitos mínimos del sistema.

    Args:
        df (pd.DataFrame): DataFrame a validar

    Returns:
        Tuple[bool, set]: 
            - True si es válido, False si faltan columnas
            - Set de nombres de columnas faltantes (vacío si es válido)

    Examples:
        >>> valido, faltantes = validar_columnas(df)
        >>> if not valido:
        >>>     print(f"Faltan columnas: {faltantes}")
    """
    try:
        safe_validate_schema(df)
        return True, set()
    except ValidationError as e:
        # Extraer nombres de columnas del mensaje de error
        faltantes = set(str(e).replace('Faltan columnas requeridas: ', '').split(', ')) if str(e) else set()
        logger.warning(f"Validación de esquema falló: {faltantes}")
        return False, faltantes


def calcular_metricas_resumen(df: pd.DataFrame) -> dict:
    """
    Calcula métricas ejecutivas principales del DataFrame con manejo robusto de errores.

    Wrapper seguro para la función resumen_metricas que incluye logging
    y manejo de excepciones para evitar que errores de cálculo rompan la aplicación.

    Args:
        df (pd.DataFrame): DataFrame con datos de ventas validados

    Returns:
        dict: Diccionario con métricas calculadas:
              - ingresos_totales: float
              - isv_total: float  
              - utilidad_promedio: float
              - unidades_vendidas: int
              - num_categorias: int
              - meses_activos: int
              Retorna diccionario vacío en caso de error.

    Examples:
        >>> metricas = calcular_metricas_resumen(df)
        >>> if metricas:
        >>>     print(f"Ingresos totales: ${metricas['ingresos_totales']:,.2f}")
    """
    try:
        metricas = resumen_metricas(df)
        logger.info(f"Métricas calculadas exitosamente: {len(metricas)} métricas")
        return metricas
    except Exception as e:
        logger.exception("Fallo calculando métricas de resumen")
        st.error(f"Error calculando métricas: {str(e)}")
        return {}

# ------------------------------
# Layout principal mejorado
# ------------------------------

# Header principal estilizado
st.markdown("""
<div class="main-header">
    <h1 style="margin: 0; font-size: 2.5rem;">📊 Sistema de Análisis de Ventas</h1>
    <h2 style="margin: 0.5rem 0 0 0; font-size: 1.5rem; opacity: 0.9;">Empresa Zambranos - Dashboard Interactivo</h2>
    <p style="margin: 0.5rem 0 0 0; opacity: 0.8;">Ciencia de Datos Fase 2 - Análisis Web Avanzado</p>
</div>
""", unsafe_allow_html=True)

# Sidebar mejorada con mejor organización
with st.sidebar:
    st.markdown("### 🔧 **Panel de Control**")
    
    # Sección de carga de datos
    with st.expander("📁 **Cargar Datos**", expanded=True):
        uploaded = st.file_uploader(
            "Selecciona tu archivo de datos",
            type=["csv", "xlsx", "xls"],
            help="Formatos soportados: CSV, Excel (.xlsx, .xls)"
        )
        
        # Guardar de forma persistente el archivo subido
        if uploaded is not None:
            content = uploaded.getvalue()
            saved_path = save_upload_safely(uploaded.name, content)
            st.success(f"✅ Archivo guardado: {saved_path.name}")
            st.session_state["last_upload"] = str(saved_path)
        
        delimiter = st.selectbox(
            "Delimitador CSV",
            ["Auto", "Coma (,)", "Punto y coma (;)", "Tab (\t)"],
            help="Solo aplica para archivos CSV"
        )

        # Selector de archivos recientes
        recent_files = sorted([p.name for p in UPLOADS_DIR.glob('*') if p.is_file() and not p.name.endswith('.lock')])
        selected_recent = None
        if recent_files:
            selected_recent = st.selectbox("Archivos recientes (uploads)", ["— ninguno —"] + recent_files)
            if selected_recent and selected_recent != "— ninguno —":
                st.info(f"Usando archivo reciente: {selected_recent}")
                # Cargar este archivo en lugar del upload directo
                uploaded = None
                class _Dummy:
                    def __init__(self, path):
                        self.name = path.name
                        self._path = path
                    def getvalue(self):
                        return open(self._path, 'rb').read()
                uploaded = _Dummy(UPLOADS_DIR / selected_recent)
    
    # Sección de configuración
    with st.expander("⚙️ **Configuración de Análisis**", expanded=True):
        k_clusters = st.slider(
            "Número de clústeres",
            min_value=2, max_value=8, value=3, step=1,
            help="Configura el número de grupos para el análisis de clustering"
        )
        
        tema_graficas = st.selectbox(
            "Tema de gráficas",
            ["Moderno", "Clásico", "Oscuro"],
            help="Selecciona el estilo visual de las gráficas"
        )
        
        mostrar_interactivas = st.checkbox(
            "Gráficas interactivas",
            value=True,
            help="Usa Plotly para gráficas más dinámicas"
        )
    
    # Sección de ayuda
    with st.expander("❓ **Ayuda y Guía**"):
        st.markdown("""
        **📋 Pasos para usar la aplicación:**
        1. Sube tu archivo de datos CSV/Excel
        2. Configura los parámetros de análisis
        3. Explora las diferentes secciones
        4. Descarga los resultados
        
        **📊 Análisis disponibles:**
        - Resumen ejecutivo con métricas clave
        - Gráficas descriptivas interactivas
        - Análisis de tendencias temporales
        - Clustering y segmentación
        - Análisis de componentes principales (PCA)
        """)
    
    st.markdown("---")
    st.markdown("💡 **Tip:** Mantén este panel abierto para acceso rápido a controles")
    
    # Tour interactivo
    mostrar_tour_interactivo()

# Carga de datos
try:
    _df, origen = load_dataframe(uploaded, uploaded.name if uploaded else None, delimiter)
except AppError as e:
    st.error(f"Error de carga de datos: {e}")
    if getattr(e, 'detail', None):
        st.caption(e.detail)
    st.stop()
if _df is None:
    st.info("👆 **Para comenzar, sube un archivo de datos usando el panel lateral**")
    st.markdown("""
    ### 📝 **Formato de datos requerido:**
    
    Tu archivo debe contener las siguientes columnas:
    - **Mes**: Nombre del mes (ej: Enero, Febrero...)
    - **Categoría**: Categoría del producto
    - **Cantidad Vendida**: Unidades vendidas
    - **Ingreso Total**: Ingresos totales
    - **ISV**: Impuesto sobre ventas
    - **Utilidad Bruta**: Ganancia bruta
    """)
    st.stop()

# Validación de datos con feedback mejorado
ok, cols_faltantes = validar_columnas(_df)
if not ok:
    st.error(f"❌ **Columnas faltantes:** {', '.join(sorted(list(cols_faltantes)))}")
    st.info("💡 **Sugerencia:** Verifica que tu archivo contenga todas las columnas requeridas")
    st.stop()

# Preprocesamiento de datos
_df_num = safe_coerce_numeric(_df.copy(), VARIABLES_NUM)

# Notificación de éxito
crear_alerta_personalizada(f"✅ **Datos cargados exitosamente** desde {origen}", "success")

# ------------------------------
# Dashboard con pestañas organizadas
# ------------------------------

# Cálculo de métricas para el resumen
metricas = calcular_metricas_resumen(_df_num)

# Pestañas principales para mejor organización
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 **Resumen Ejecutivo**", 
    "📈 **Análisis Temporal**", 
    "🎯 **Análisis por Categoría**", 
    "🔍 **Análisis Estadístico**", 
    "🧮 **Machine Learning**"
])

with tab1:
    mostrar_header_seccion("Resumen Ejecutivo", "Métricas clave y KPIs del negocio")
    mostrar_ayuda_contextual("resumen")
    
    # Métricas principales en tarjetas
    if metricas:
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            crear_tarjeta_metrica(
                "💰 Ingresos Totales",
                f"L {metricas['total_ingresos']:,.2f}" if 'total_ingresos' in metricas else "N/A"
            )
        
        with col2:
            crear_tarjeta_metrica(
                "🏛️ ISV Total",
                f"L {metricas['total_isv']:,.2f}" if 'total_isv' in metricas else "N/A"
            )
        
        with col3:
            crear_tarjeta_metrica(
                "📈 Utilidad Promedio",
                f"L {metricas['utilidad_promedio']:,.2f}" if 'utilidad_promedio' in metricas else "N/A"
            )
        
        with col4:
            crear_tarjeta_metrica(
                "🛒 Unidades Vendidas",
                f"{metricas['total_ventas']:,}" if 'total_ventas' in metricas else "N/A"
            )
        
        st.markdown("---")
        
        # Métricas secundarias
        col5, col6, col7 = st.columns(3)
        
        with col5:
            crear_tarjeta_metrica(
                "🏪 Categorías",
                f"{metricas['categorias_unicas']}" if 'categorias_unicas' in metricas else "N/A"
            )
        
        with col6:
            crear_tarjeta_metrica(
                "📅 Meses Activos",
                f"{metricas['meses_activos']}" if 'meses_activos' in metricas else "N/A"
            )
        
        with col7:
            ingreso_promedio_mes = metricas.get('total_ingresos', 0) / max(metricas.get('meses_activos', 1), 1)
            crear_tarjeta_metrica(
                "📊 Ingreso Promedio/Mes",
                f"L {ingreso_promedio_mes:,.2f}"
            )
    
    # Vista previa de datos con filtros
    mostrar_header_seccion("Vista de Datos", "Exploración interactiva del dataset")
    
    col1, col2 = st.columns([2, 1])
    
    with col2:
        filtro_mes = st.selectbox(
            "Filtrar por mes:",
            ["Todos"] + MESES_ORDEN,
            help="Filtra los datos por mes específico"
        )
        
        filtro_categoria = st.selectbox(
            "Filtrar por categoría:",
            ["Todas"] + sorted(_df["Categoría"].unique().tolist()),
            help="Filtra los datos por categoría"
        )
        
        mostrar_filas = st.number_input(
            "Filas a mostrar:",
            min_value=5, max_value=100, value=10, step=5
        )
    
    with col1:
        # Aplicar filtros
        df_filtrado = _df_num.copy()
        if filtro_mes != "Todos":
            df_filtrado = df_filtrado[df_filtrado["Mes"] == filtro_mes]
        if filtro_categoria != "Todas":
            df_filtrado = df_filtrado[df_filtrado["Categoría"] == filtro_categoria]
        
        st.dataframe(
            df_filtrado.head(mostrar_filas),
            use_container_width=True,
            height=400
        )
        
        st.caption(f"Mostrando {min(mostrar_filas, len(df_filtrado))} de {len(df_filtrado)} registros filtrados")

    # Exportaciones
    mostrar_header_seccion("Exportaciones", "Descarga de datos y reporte")
    colx, coly, colz = st.columns(3)
    with colx:
        exportar_dataframe_csv(_df_num, "datos_filtrados" if (filtro_mes!="Todos" or filtro_categoria!="Todas") else "datos_completos")
    with coly:
        metricas_doc = generar_reporte_docx(metricas, notas="Reporte generado desde la app web")
        st.download_button(
            label="📝 Descargar Reporte DOCX",
            data=metricas_doc,
            file_name="reporte_ventas.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with colz:
        # ZIP simple con CSV original si existe en memoria
        csv_bytes = _df_num.to_csv(index=False).encode('utf-8')
        exportar_zip({"datos.csv": csv_bytes}, "paquete_datos.zip")

with tab2:
    mostrar_header_seccion("Análisis Temporal", "Tendencias y patrones a lo largo del tiempo")
    mostrar_ayuda_contextual("temporal")
    
    # Selector de análisis temporal
    analisis_temporal = st.radio(
        "Selecciona el tipo de análisis:",
        ["📈 Ingresos por Mes", "🏛️ ISV por Mes", "📊 Comparativo Mensual"],
        horizontal=True
    )
    
    if analisis_temporal == "📈 Ingresos por Mes":
        try:
            ingresos_mensuales = _df_num.groupby("Mes")["Ingreso Total"].sum().reindex(MESES_ORDEN)
            
            if mostrar_interactivas:
                # Gráfica interactiva con Plotly
                fig = px.bar(
                    x=ingresos_mensuales.index,
                    y=ingresos_mensuales.values,
                    title="Ingresos Totales por Mes",
                    labels={"x": "Mes", "y": "Ingresos (Lempiras)"},
                    color=ingresos_mensuales.values,
                    color_continuous_scale="Blues"
                )
                fig.update_layout(height=500, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
                _record_plotly(fig, "ingresos_mensuales")
            else:
                # Gráfica estática mejorada
                fig, ax = plt.subplots(figsize=(12, 6))
                bars = ingresos_mensuales.plot(kind="bar", color="skyblue", ax=ax)
                ax.set_title("Ingresos Totales por Mes", fontsize=16, fontweight='bold')
                ax.set_ylabel("Lempiras", fontsize=12)
                ax.set_xlabel("Mes", fontsize=12)
                ax.tick_params(axis='x', rotation=45)
                ax.grid(axis='y', alpha=0.3)
                
                # Añadir valores en las barras
                for bar in bars.patches:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'L{height:,.0f}', ha='center', va='bottom', fontsize=10)
                
                plt.tight_layout()
                st.pyplot(fig)
                _record_matplotlib(fig, "ingresos_mensuales")
                
                # Botón de descarga
                st.markdown(generar_download_link(fig, "ingresos_mensuales"), unsafe_allow_html=True)
            
            # Análisis adicional
            max_mes = ingresos_mensuales.idxmax()
            min_mes = ingresos_mensuales.idxmin()
            st.info(f"🏆 **Mejor mes:** {max_mes} (L{ingresos_mensuales[max_mes]:,.2f})")
            st.warning(f"📉 **Mes más bajo:** {min_mes} (L{ingresos_mensuales[min_mes]:,.2f})")
            
        except Exception as e:
            st.error(f"Error al generar la gráfica de ingresos: {e}")
    
    elif analisis_temporal == "🏛️ ISV por Mes":
        try:
            isv_mensual = _df_num.groupby("Mes")["ISV"].sum().reindex(MESES_ORDEN)
            
            if mostrar_interactivas:
                fig = px.line(
                    x=isv_mensual.index,
                    y=isv_mensual.values,
                    title="ISV Generado por Mes",
                    labels={"x": "Mes", "y": "ISV (Lempiras)"},
                    markers=True
                )
                fig.update_traces(line_color='green', line_width=3, marker_size=8)
                fig.update_layout(height=500)
                st.plotly_chart(fig, use_container_width=True)
                _record_plotly(fig, "isv_mensual")
            else:
                fig, ax = plt.subplots(figsize=(12, 6))
                isv_mensual.plot(kind="line", marker='o', color='green', linewidth=2, markersize=8, ax=ax)
                ax.set_title("ISV Generado por Mes", fontsize=16, fontweight='bold')
                ax.set_ylabel("ISV (Lempiras)", fontsize=12)
                ax.set_xlabel("Mes", fontsize=12)
                ax.grid(True, alpha=0.3)
                ax.tick_params(axis='x', rotation=45)
                plt.tight_layout()
                st.pyplot(fig)
                _record_matplotlib(fig, "isv_mensual")
                
                st.markdown(generar_download_link(fig, "isv_mensual"), unsafe_allow_html=True)
            
            # Estadísticas del ISV
            total_isv = isv_mensual.sum()
            promedio_isv = isv_mensual.mean()
            st.success(f"💰 **ISV Total del año:** L{total_isv:,.2f}")
            st.info(f"📊 **ISV Promedio mensual:** L{promedio_isv:,.2f}")
            
        except Exception as e:
            st.error(f"Error al generar la gráfica de ISV: {e}")
    
    else:  # Comparativo Mensual
        try:
            # Datos para comparativo
            ingresos_mensuales = _df_num.groupby("Mes")["Ingreso Total"].sum().reindex(MESES_ORDEN)
            isv_mensual = _df_num.groupby("Mes")["ISV"].sum().reindex(MESES_ORDEN)
            
            if mostrar_interactivas:
                # Gráfica comparativa con dos ejes Y
                fig = make_subplots(specs=[[{"secondary_y": True}]])
                
                fig.add_trace(
                    go.Bar(x=ingresos_mensuales.index, y=ingresos_mensuales.values, 
                          name="Ingresos", marker_color="lightblue"),
                    secondary_y=False,
                )
                
                fig.add_trace(
                    go.Scatter(x=isv_mensual.index, y=isv_mensual.values, 
                              name="ISV", mode="lines+markers", line_color="green"),
                    secondary_y=True,
                )
                
                fig.update_xaxes(title_text="Mes")
                fig.update_yaxes(title_text="Ingresos (Lempiras)", secondary_y=False)
                fig.update_yaxes(title_text="ISV (Lempiras)", secondary_y=True)
                fig.update_layout(title_text="Comparativo: Ingresos vs ISV por Mes", height=500)
                
                st.plotly_chart(fig, use_container_width=True)
                _record_plotly(fig, "comparativo_mensual")
            else:
                fig, ax1 = plt.subplots(figsize=(12, 6))
                
                # Ingresos (barras)
                ax1.bar(ingresos_mensuales.index, ingresos_mensuales.values, 
                       alpha=0.7, color='lightblue', label='Ingresos')
                ax1.set_xlabel('Mes')
                ax1.set_ylabel('Ingresos (Lempiras)', color='blue')
                ax1.tick_params(axis='x', rotation=45)
                
                # ISV (línea)
                ax2 = ax1.twinx()
                ax2.plot(isv_mensual.index, isv_mensual.values, 
                        color='green', marker='o', linewidth=2, markersize=6, label='ISV')
                ax2.set_ylabel('ISV (Lempiras)', color='green')
                
                plt.title('Comparativo: Ingresos vs ISV por Mes', fontsize=16, fontweight='bold')
                
                # Leyendas
                lines1, labels1 = ax1.get_legend_handles_labels()
                lines2, labels2 = ax2.get_legend_handles_labels()
                ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left')
                
                plt.tight_layout()
                st.pyplot(fig)
                _record_matplotlib(fig, "comparativo_mensual")
                
                st.markdown(generar_download_link(fig, "comparativo_mensual"), unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"Error al generar el comparativo mensual: {e}")

    # Exportación global de figuras del análisis temporal
    if st.button("📦 Descargar todas las gráficas (Temporal) como ZIP"):
        if st.session_state["_assets"]:
            exportar_zip(st.session_state["_assets"], "graficas_temporal.zip")
        else:
            st.info("No hay gráficas generadas aún.")

with tab3:
    mostrar_header_seccion("Análisis por Categoría", "Rendimiento y distribución por categorías de productos")
    
    # Selectores para análisis por categoría
    analisis_categoria = st.radio(
        "Selecciona el análisis:",
        ["🥧 Distribución de Ventas", "💰 Utilidad por Categoría", "📊 Ranking de Categorías"],
        horizontal=True
    )
    
    if analisis_categoria == "🥧 Distribución de Ventas":
        try:
            ventas_categoria = _df_num.groupby("Categoría")["Cantidad Vendida"].sum()
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                if mostrar_interactivas:
                    fig = px.pie(
                        values=ventas_categoria.values,
                        names=ventas_categoria.index,
                        title="Distribución de Ventas por Categoría",
                        color_discrete_sequence=px.colors.qualitative.Set3
                    )
                    fig.update_traces(textposition='inside', textinfo='percent+label')
                    fig.update_layout(height=500)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    fig, ax = plt.subplots(figsize=(10, 8))
                    colors = plt.cm.Set3(np.linspace(0, 1, len(ventas_categoria)))
                    wedges, texts, autotexts = ax.pie(
                        ventas_categoria.values, 
                        labels=ventas_categoria.index,
                        autopct='%1.1f%%', 
                        startangle=90,
                        colors=colors
                    )
                    ax.set_title("Distribución de Ventas por Categoría", fontsize=16, fontweight='bold')
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    st.markdown(generar_download_link(fig, "distribucion_ventas"), unsafe_allow_html=True)
            
            with col2:
                st.markdown("#### 📈 **Estadísticas por Categoría**")
                for categoria, cantidad in ventas_categoria.sort_values(ascending=False).items():
                    porcentaje = (cantidad / ventas_categoria.sum()) * 100
                    st.metric(
                        f"**{categoria}**",
                        f"{cantidad:,} unidades",
                        f"{porcentaje:.1f}% del total"
                    )
                # Exportación CSV de ventas por categoría
                exportar_dataframe_csv(ventas_categoria.reset_index().rename(columns={"Cantidad Vendida":"Unidades"}), "ventas_por_categoria")
                    
        except Exception as e:
            st.error(f"Error al generar la distribución de ventas: {e}")
    
    elif analisis_categoria == "💰 Utilidad por Categoría":
        try:
            utilidad_por_categoria = _df_num.groupby("Categoría")["Utilidad Bruta"].agg(['mean', 'sum']).round(2)
            utilidad_por_categoria.columns = ['Promedio', 'Total']
            
            # Selector de métrica
            metrica_utilidad = st.selectbox(
                "Métrica a mostrar:",
                ["Promedio", "Total"],
                help="Selecciona si quieres ver la utilidad promedio o total por categoría"
            )
            
            data_to_plot = utilidad_por_categoria[metrica_utilidad].sort_values()
            
            if mostrar_interactivas:
                fig = px.bar(
                    x=data_to_plot.values,
                    y=data_to_plot.index,
                    orientation='h',
                    title=f"Utilidad Bruta {metrica_utilidad} por Categoría",
                    labels={"x": "Utilidad (Lempiras)", "y": "Categoría"},
                    color=data_to_plot.values,
                    color_continuous_scale="Oranges"
                )
                fig.update_layout(height=500, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
                _record_plotly(fig, "ventas_categoria_pie")
            else:
                fig, ax = plt.subplots(figsize=(12, 6))
                bars = data_to_plot.plot(kind="barh", color="orange", ax=ax)
                ax.set_title(f"Utilidad Bruta {metrica_utilidad} por Categoría", fontsize=16, fontweight='bold')
                ax.set_xlabel("Utilidad (Lempiras)", fontsize=12)
                ax.grid(axis='x', alpha=0.3)
                
                # Añadir valores en las barras
                for i, bar in enumerate(bars.patches):
                    width = bar.get_width()
                    ax.text(width, bar.get_y() + bar.get_height()/2.,
                           f'L{width:,.0f}', ha='left', va='center', fontsize=10)
                
                plt.tight_layout()
                st.pyplot(fig)
                _record_matplotlib(fig, "ventas_categoria_pie")
                
                st.markdown(generar_download_link(fig, f"utilidad_{metrica_utilidad.lower()}"), unsafe_allow_html=True)
            
            # Tabla resumen
            st.markdown("#### 📋 **Resumen de Utilidades por Categoría**")
            st.dataframe(
                utilidad_por_categoria.style.format("L{:,.2f}"),
                use_container_width=True
            )
            exportar_dataframe_csv(utilidad_por_categoria.reset_index(), "utilidad_por_categoria")
            
        except Exception as e:
            st.error(f"Error al generar el análisis de utilidad: {e}")
    
    else:  # Ranking de Categorías
        try:
            # Métricas múltiples por categoría
            ranking_data = _df_num.groupby("Categoría").agg({
                "Cantidad Vendida": "sum",
                "Ingreso Total": "sum",
                "Utilidad Bruta": "mean",
                "ISV": "sum"
            }).round(2)
            
            ranking_data.columns = ["Unidades Vendidas", "Ingresos Totales", "Utilidad Promedio", "ISV Total"]
            
            # Selector de métrica para ranking
            metrica_ranking = st.selectbox(
                "Ordenar ranking por:",
                ["Ingresos Totales", "Unidades Vendidas", "Utilidad Promedio", "ISV Total"]
            )
            
            ranking_ordenado = ranking_data.sort_values(metrica_ranking, ascending=False)
            
            # Mostrar ranking
            st.markdown(f"#### 🏆 **Ranking por {metrica_ranking}**")
            
            # Top 3 destacado
            col1, col2, col3 = st.columns(3)
            top_3 = ranking_ordenado.head(3)
            
            with col1:
                if len(top_3) >= 1:
                    categoria = top_3.index[0]
                    valor = top_3.iloc[0][metrica_ranking]
                    st.markdown(f"""
                    <div style="background: linear-gradient(45deg, #FFD700, #FFA500); padding: 1rem; border-radius: 10px; text-align: center;">
                        <h3>🥇 #{1}</h3>
                        <h4>{categoria}</h4>
                        <p><strong>L{valor:,.2f}</strong></p>
                    </div>
                    """, unsafe_allow_html=True)
            
            with col2:
                if len(top_3) >= 2:
                    categoria = top_3.index[1]
                    valor = top_3.iloc[1][metrica_ranking]
                    st.markdown(f"""
                    <div style="background: linear-gradient(45deg, #C0C0C0, #A9A9A9); padding: 1rem; border-radius: 10px; text-align: center;">
                        <h3>🥈 #{2}</h3>
                        <h4>{categoria}</h4>
                        <p><strong>L{valor:,.2f}</strong></p>
                    </div>
                    """, unsafe_allow_html=True)
            
            with col3:
                if len(top_3) >= 3:
                    categoria = top_3.index[2]
                    valor = top_3.iloc[2][metrica_ranking]
                    st.markdown(f"""
                    <div style="background: linear-gradient(45deg, #CD7F32, #8B4513); padding: 1rem; border-radius: 10px; text-align: center;">
                        <h3>🥉 #{3}</h3>
                        <h4>{categoria}</h4>
                        <p><strong>L{valor:,.2f}</strong></p>
                    </div>
                    """, unsafe_allow_html=True)
            
            st.markdown("---")
            
            # Tabla completa del ranking
            st.markdown("#### 📊 **Ranking Completo**")
            
            # Añadir posiciones al dataframe
            ranking_con_posicion = ranking_ordenado.copy()
            ranking_con_posicion.insert(0, "Posición", range(1, len(ranking_con_posicion) + 1))
            
            st.dataframe(
                ranking_con_posicion.style.format({
                    "Unidades Vendidas": "{:,}",
                    "Ingresos Totales": "L{:,.2f}",
                    "Utilidad Promedio": "L{:,.2f}",
                    "ISV Total": "L{:,.2f}"
                }),
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"Error al generar el ranking de categorías: {e}")

    if st.button("📦 Descargar todas las gráficas (Categoría) como ZIP"):
        if st.session_state["_assets"]:
            exportar_zip(st.session_state["_assets"], "graficas_categoria.zip")
        else:
            st.info("No hay gráficas generadas aún.")

with tab4:
    mostrar_header_seccion("Análisis Estadístico", "Distribuciones y patrones estadísticos en los datos")
    
    # Selector de análisis estadístico
    analisis_estadistico = st.radio(
        "Tipo de análisis:",
        ["📊 Histogramas", "📦 Boxplots", "🔗 Correlaciones", "📈 Estadísticas Descriptivas"],
        horizontal=True
    )
    
    if analisis_estadistico == "📊 Histogramas":
        try:
            vars_existentes = [c for c in VARIABLES_NUM if c in _df_num.columns]
            if len(vars_existentes) == 0:
                st.info("No hay variables numéricas disponibles.")
            else:
                # Selector de variables
                variables_seleccionadas = st.multiselect(
                    "Selecciona las variables a analizar:",
                    vars_existentes,
                    default=vars_existentes[:4]  # Primeras 4 por defecto
                )
                
                if variables_seleccionadas:
                    if mostrar_interactivas:
                        # Crear subplots con Plotly
                        n_vars = len(variables_seleccionadas)
                        n_cols = min(2, n_vars)
                        n_rows = (n_vars + n_cols - 1) // n_cols
                        
                        fig = make_subplots(
                            rows=n_rows, cols=n_cols,
                            subplot_titles=variables_seleccionadas
                        )
                        
                        for i, var in enumerate(variables_seleccionadas):
                            row = i // n_cols + 1
                            col = i % n_cols + 1
                            
                            fig.add_trace(
                                go.Histogram(x=_df_num[var].dropna(), name=var, showlegend=False),
                                row=row, col=col
                            )
                        
                        fig.update_layout(height=300 * n_rows, title_text="Distribución de Variables")
                        st.plotly_chart(fig, use_container_width=True)
                        _record_plotly(fig, "utilidad_categoria")
                    else:
                        n_vars = len(variables_seleccionadas)
                        n_cols = min(3, n_vars)
                        n_rows = (n_vars + n_cols - 1) // n_cols
                        
                        fig, axes = plt.subplots(n_rows, n_cols, figsize=(15, 4 * n_rows))
                        if n_vars == 1:
                            axes = [axes]
                        elif n_rows == 1:
                            axes = axes.reshape(1, -1)
                        
                        for i, var in enumerate(variables_seleccionadas):
                            row, col = i // n_cols, i % n_cols
                            ax = axes[row, col] if n_rows > 1 else axes[col]
                            
                            _df_num[var].hist(bins=15, edgecolor='black', alpha=0.7, ax=ax)
                            ax.set_title(f'Distribución de {var}', fontweight='bold')
                            ax.set_xlabel(var)
                            ax.set_ylabel('Frecuencia')
                            ax.grid(alpha=0.3)
                        
                        # Ocultar subplots vacíos
                        for j in range(len(variables_seleccionadas), n_rows * n_cols):
                            row, col = j // n_cols, j % n_cols
                            ax = axes[row, col] if n_rows > 1 else axes[col]
                            ax.axis('off')
                        
                        plt.tight_layout()
                        st.pyplot(fig)
                        _record_matplotlib(fig, "utilidad_categoria")
                        
                        st.markdown(generar_download_link(fig, "histogramas"), unsafe_allow_html=True)
                
        except Exception as e:
            st.error(f"Error al generar histogramas: {e}")
    
    elif analisis_estadistico == "📦 Boxplots":
        try:
            vars_existentes = [c for c in VARIABLES_NUM if c in _df_num.columns]
            if len(vars_existentes) == 0:
                st.info("No hay variables numéricas disponibles.")
            else:
                variables_seleccionadas = st.multiselect(
                    "Selecciona las variables para boxplot:",
                    vars_existentes,
                    default=vars_existentes[:6]
                )
                
                if variables_seleccionadas:
                    if mostrar_interactivas:
                        fig = go.Figure()
                        for var in variables_seleccionadas:
                            fig.add_trace(go.Box(y=_df_num[var].dropna(), name=var))
                        
                        fig.update_layout(
                            title="Análisis de Outliers y Distribución",
                            yaxis_title="Valores",
                            height=500
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        fig, ax = plt.subplots(figsize=(12, 6))
                        data_for_boxplot = [_df_num[var].dropna() for var in variables_seleccionadas]
                        
                        bp = ax.boxplot(data_for_boxplot, labels=variables_seleccionadas, patch_artist=True)
                        
                        # Colorear las cajas
                        colors = plt.cm.Set3(np.linspace(0, 1, len(variables_seleccionadas)))
                        for patch, color in zip(bp['boxes'], colors):
                            patch.set_facecolor(color)
                        
                        ax.set_title("Análisis de Outliers y Distribución", fontsize=16, fontweight='bold')
                        ax.tick_params(axis='x', rotation=45)
                        ax.grid(alpha=0.3)
                        plt.tight_layout()
                        st.pyplot(fig)
                        
                        st.markdown(generar_download_link(fig, "boxplots"), unsafe_allow_html=True)
                    
                    # Análisis de outliers
                    st.markdown("#### 🔍 **Análisis de Outliers**")
                    for var in variables_seleccionadas:
                        Q1 = _df_num[var].quantile(0.25)
                        Q3 = _df_num[var].quantile(0.75)
                        IQR = Q3 - Q1
                        outliers = _df_num[((_df_num[var] < (Q1 - 1.5 * IQR)) | (_df_num[var] > (Q3 + 1.5 * IQR)))]
                        st.write(f"**{var}:** {len(outliers)} outliers detectados ({len(outliers)/len(_df_num)*100:.1f}% del total)")
                
        except Exception as e:
            st.error(f"Error al generar boxplots: {e}")
    
    elif analisis_estadistico == "🔗 Correlaciones":
        try:
            vars_existentes = [c for c in VARIABLES_NUM if c in _df_num.columns]
            if len(vars_existentes) < 2:
                st.info("Se necesitan al menos 2 variables numéricas para análisis de correlación.")
            else:
                # Matrix de correlación
                corr_matrix = _df_num[vars_existentes].corr()
                
                if mostrar_interactivas:
                    fig = px.imshow(
                        corr_matrix,
                        text_auto=True,
                        aspect="auto",
                        title="Matriz de Correlación",
                        color_continuous_scale="RdBu"
                    )
                    fig.update_layout(height=600)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    fig, ax = plt.subplots(figsize=(12, 8))
                    mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
                    sns.heatmap(
                        corr_matrix, 
                        mask=mask,
                        annot=True, 
                        cmap="RdBu", 
                        center=0,
                        square=True,
                        fmt='.2f',
                        ax=ax
                    )
                    ax.set_title("Matriz de Correlación", fontsize=16, fontweight='bold')
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    st.markdown(generar_download_link(fig, "correlaciones"), unsafe_allow_html=True)
                
                # Correlaciones más fuertes
                st.markdown("#### 🔗 **Correlaciones más Fuertes**")
                
                # Obtener correlaciones sin la diagonal
                corr_pairs = []
                for i in range(len(corr_matrix.columns)):
                    for j in range(i+1, len(corr_matrix.columns)):
                        var1 = corr_matrix.columns[i]
                        var2 = corr_matrix.columns[j]
                        corr_val = corr_matrix.iloc[i, j]
                        corr_pairs.append((var1, var2, corr_val))
                
                # Ordenar por valor absoluto de correlación
                corr_pairs.sort(key=lambda x: abs(x[2]), reverse=True)
                
                for var1, var2, corr_val in corr_pairs[:5]:  # Top 5
                    if corr_val > 0.7:
                        st.success(f"🔗 **{var1}** ↔ **{var2}**: {corr_val:.3f} (Correlación fuerte positiva)")
                    elif corr_val < -0.7:
                        st.error(f"🔗 **{var1}** ↔ **{var2}**: {corr_val:.3f} (Correlación fuerte negativa)")
                    elif abs(corr_val) > 0.3:
                        st.info(f"🔗 **{var1}** ↔ **{var2}**: {corr_val:.3f} (Correlación moderada)")
                    else:
                        st.write(f"🔗 **{var1}** ↔ **{var2}**: {corr_val:.3f} (Correlación débil)")
                
        except Exception as e:
            st.error(f"Error al generar análisis de correlaciones: {e}")
    
    else:  # Estadísticas Descriptivas
        try:
            vars_existentes = [c for c in VARIABLES_NUM if c in _df_num.columns]
            if len(vars_existentes) == 0:
                st.info("No hay variables numéricas disponibles.")
            else:
                # Estadísticas descriptivas completas
                stats_desc = _df_num[vars_existentes].describe().round(2)
                
                st.markdown("#### 📈 **Estadísticas Descriptivas Completas**")
                st.dataframe(
                    stats_desc.style.format("{:,.2f}"),
                    use_container_width=True
                )
                
                # Información adicional
                st.markdown("#### ℹ️ **Información Adicional**")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**📊 Datos Faltantes:**")
                    missing_data = _df_num[vars_existentes].isnull().sum()
                    for var, missing in missing_data.items():
                        if missing > 0:
                            st.warning(f"**{var}:** {missing} valores faltantes ({missing/len(_df_num)*100:.1f}%)")
                        else:
                            st.success(f"**{var}:** Sin datos faltantes ✅")
                
                with col2:
                    st.markdown("**📏 Rangos de Datos:**")
                    for var in vars_existentes:
                        min_val = _df_num[var].min()
                        max_val = _df_num[var].max()
                        st.info(f"**{var}:** L{min_val:,.2f} - L{max_val:,.2f}")
                
        except Exception as e:
            st.error(f"Error al generar estadísticas descriptivas: {e}")

with tab5:
    mostrar_header_seccion("Machine Learning", "Análisis avanzado con PCA y Clustering")
    
    # Selector de análisis ML
    analisis_ml = st.radio(
        "Tipo de análisis:",
        ["🔍 Análisis PCA", "🎯 Clustering", "📊 Análisis Combinado"],
        horizontal=True
    )
    
    vars_existentes = [c for c in VARIABLES_NUM if c in _df_num.columns]
    X = _df_num[vars_existentes].dropna()
    
    if len(vars_existentes) < 2 or X.shape[0] < k_clusters:
        st.warning("⚠️ Datos insuficientes para análisis de Machine Learning.")
        st.info("Se requieren al menos 2 variables numéricas y datos suficientes para el número de clústeres seleccionado.")
    else:
        if analisis_ml == "🔍 Análisis PCA":
            try:
                # Configuración PCA
                col1, col2 = st.columns([1, 3])
                
                with col1:
                    n_components = st.slider(
                        "Componentes PCA:",
                        min_value=2, max_value=min(len(vars_existentes), 5), value=2
                    )
                    
                    mostrar_varianza = st.checkbox("Mostrar varianza explicada", value=True)
                
                with col2:
                    # Ejecutar PCA
                    X_scaled = StandardScaler().fit_transform(X)
                    pca = PCA(n_components=n_components)
                    pca_result = pca.fit_transform(X_scaled)
                    
                    if mostrar_interactivas and n_components >= 2:
                        fig = px.scatter(
                            x=pca_result[:, 0], y=pca_result[:, 1],
                            title="Análisis de Componentes Principales (PCA)",
                            labels={"x": "Componente Principal 1", "y": "Componente Principal 2"}
                        )
                        fig.update_layout(height=500)
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        fig, ax = plt.subplots(figsize=(10, 6))
                        scatter = ax.scatter(pca_result[:, 0], pca_result[:, 1], 
                                           alpha=0.7, s=50, c=range(len(pca_result)), cmap='viridis')
                        ax.set_xlabel(f"Componente Principal 1 ({pca.explained_variance_ratio_[0]:.1%} varianza)")
                        ax.set_ylabel(f"Componente Principal 2 ({pca.explained_variance_ratio_[1]:.1%} varianza)")
                        ax.set_title("Análisis de Componentes Principales (PCA)")
                        ax.grid(alpha=0.3)
                        plt.colorbar(scatter, label='Punto de datos')
                        plt.tight_layout()
                        st.pyplot(fig)
                        
                        st.markdown(generar_download_link(fig, "pca_analysis"), unsafe_allow_html=True)
                
                # Información de varianza explicada
                if mostrar_varianza:
                    st.markdown("#### 📊 **Varianza Explicada por Componente**")
                    
                    var_exp_df = pd.DataFrame({
                        'Componente': [f'PC{i+1}' for i in range(n_components)],
                        'Varianza Individual': pca.explained_variance_ratio_,
                        'Varianza Acumulada': np.cumsum(pca.explained_variance_ratio_)
                    })
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.dataframe(
                            var_exp_df.style.format({
                                'Varianza Individual': '{:.1%}',
                                'Varianza Acumulada': '{:.1%}'
                            }),
                            use_container_width=True
                        )
                    
                    with col2:
                        if mostrar_interactivas:
                            fig = px.bar(
                                var_exp_df, x='Componente', y='Varianza Individual',
                                title="Varianza Explicada por Componente"
                            )
                            st.plotly_chart(fig, use_container_width=True)
                        
            except Exception as e:
                st.error(f"Error en análisis PCA: {e}")
        
        elif analisis_ml == "🎯 Clustering":
            try:
                # Ejecutar clustering
                X_scaled = StandardScaler().fit_transform(X)
                kmeans = KMeans(n_clusters=k_clusters, random_state=42)
                clusters = kmeans.fit_predict(X_scaled)
                
                # Añadir clusters al dataframe temporal
                df_tmp = _df_num.loc[X.index].copy()
                df_tmp['Cluster'] = clusters
                
                # Visualización de clusters en espacio PCA
                pca = PCA(n_components=2)
                pca_result = pca.fit_transform(X_scaled)
                
                if mostrar_interactivas:
                    fig = px.scatter(
                        x=pca_result[:, 0], y=pca_result[:, 1],
                        color=clusters.astype(str),
                        title=f"Clustering K-Means ({k_clusters} grupos)",
                        labels={"x": "Componente Principal 1", "y": "Componente Principal 2", "color": "Cluster"}
                    )
                    fig.update_layout(height=500)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    fig, ax = plt.subplots(figsize=(10, 6))
                    scatter = ax.scatter(pca_result[:, 0], pca_result[:, 1], 
                                       c=clusters, cmap='Set1', alpha=0.7, s=50)
                    ax.set_xlabel("Componente Principal 1")
                    ax.set_ylabel("Componente Principal 2")
                    ax.set_title(f"Clustering K-Means ({k_clusters} grupos)")
                    
                    # Añadir centroids
                    centroids_pca = pca.transform(kmeans.cluster_centers_)
                    ax.scatter(centroids_pca[:, 0], centroids_pca[:, 1], 
                             c='red', marker='x', s=200, linewidths=3, label='Centroids')
                    
                    ax.legend()
                    ax.grid(alpha=0.3)
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    st.markdown(generar_download_link(fig, "clustering"), unsafe_allow_html=True)
                
                # Perfil de clusters (heatmap)
                st.markdown("#### 🔥 **Perfil de Clusters**")
                
                perfil = df_tmp.groupby('Cluster')[vars_existentes].mean().round(2)
                
                if mostrar_interactivas:
                    fig = px.imshow(
                        perfil.values,
                        x=perfil.columns,
                        y=[f'Cluster {i}' for i in perfil.index],
                        text_auto=True,
                        aspect="auto",
                        title="Perfil Promedio por Clúster",
                        color_continuous_scale="YlOrRd"
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    fig, ax = plt.subplots(figsize=(12, 6))
                    sns.heatmap(perfil, annot=True, cmap="YlOrRd", fmt=".2f", ax=ax)
                    ax.set_title("Perfil Promedio por Clúster", fontsize=16, fontweight='bold')
                    ax.set_xlabel("Variables")
                    ax.set_ylabel("Clúster")
                    plt.tight_layout()
                    st.pyplot(fig)
                    
                    st.markdown(generar_download_link(fig, "cluster_profile"), unsafe_allow_html=True)
                        
            except Exception as e:
                st.error(f"Error en análisis de clustering: {e}")
        
        else:  # Análisis Combinado
            try:
                st.markdown("#### 🔄 **Análisis Integrado: PCA + Clustering**")
                
                # Ejecutar ambos análisis
                X_scaled = StandardScaler().fit_transform(X)
                
                # PCA
                pca = PCA(n_components=2)
                pca_result = pca.fit_transform(X_scaled)
                
                # Clustering
                kmeans = KMeans(n_clusters=k_clusters, random_state=42)
                clusters = kmeans.fit_predict(X_scaled)
                
                # Visualización combinada
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**🔍 Espacio Original (PCA)**")
                    if mostrar_interactivas:
                        fig = px.scatter(
                            x=pca_result[:, 0], y=pca_result[:, 1],
                            title="Datos en Espacio PCA",
                            labels={"x": "PC1", "y": "PC2"}
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    st.markdown("**🎯 Con Clusters Aplicados**")
                    if mostrar_interactivas:
                        fig = px.scatter(
                            x=pca_result[:, 0], y=pca_result[:, 1],
                            color=clusters.astype(str),
                            title="Clusters en Espacio PCA",
                            labels={"x": "PC1", "y": "PC2", "color": "Cluster"}
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                # Resumen del análisis combinado
                st.markdown("#### 📋 **Resumen del Análisis**")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    crear_tarjeta_metrica(
                        "🔍 Varianza Explicada (PC1+PC2)",
                        f"{(pca.explained_variance_ratio_[:2].sum()*100):.1f}%"
                    )
                
                with col2:
                    crear_tarjeta_metrica(
                        "🎯 Número de Clusters",
                        f"{k_clusters}"
                    )
                
                with col3:
                    # Calcular inercia del clustering
                    inertia = kmeans.inertia_
                    crear_tarjeta_metrica(
                        "📊 Inercia del Clustering",
                        f"{inertia:.2f}"
                    )
                
                # Recomendaciones
                st.markdown("#### 💡 **Recomendaciones**")
                
                varianza_total = pca.explained_variance_ratio_[:2].sum()
                if varianza_total > 0.8:
                    st.success("🎯 **Excelente reducción dimensional** - Los primeros 2 componentes capturan >80% de la varianza")
                elif varianza_total > 0.6:
                    st.info("✅ **Buena reducción dimensional** - Los primeros 2 componentes capturan >60% de la varianza")
                else:
                    st.warning("⚠️ **Reducción dimensional limitada** - Considera usar más componentes o revisar las variables")
                
                from sklearn.metrics import silhouette_score
                silhouette_avg = silhouette_score(X_scaled, clusters)
                
                if silhouette_avg > 0.5:
                    st.success(f"🎯 **Clusters bien definidos** - Score de silueta: {silhouette_avg:.3f}")
                else:
                    st.warning(f"⚠️ **Clusters con solapamiento** - Score de silueta: {silhouette_avg:.3f}. Considera ajustar el número de clusters.")
                    
            except Exception as e:
                st.error(f"Error en análisis combinado: {e}")

# ------------------------------
# Footer y información adicional
# ------------------------------
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem; background: linear-gradient(90deg, #f8f9fa 0%, #e9ecef 100%); border-radius: 10px; margin-top: 2rem;">
    <h4>🎓 Sistema de Análisis de Ventas - Zambranos</h4>
    <p><strong>Ciencia de Datos Fase 2</strong> | Análisis Web Interactivo</p>
    <p style="font-size: 0.9rem; color: #6c757d;">
        💡 <strong>Tip:</strong> Utiliza el panel lateral para ajustar configuraciones y explorar diferentes análisis.<br>
        📊 Esta aplicación replica y mejora los análisis de la Fase 1 con una interfaz web moderna e interactiva.
    </p>
</div>
""", unsafe_allow_html=True)
